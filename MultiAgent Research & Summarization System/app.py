import os
import streamlit as st
import pdfplumber
from docx import Document as DocxDocument

from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langgraph.graph import StateGraph
from langchain_core.runnables import RunnableLambda

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_community.tools import DuckDuckGoSearchRun
from dotenv import load_dotenv

load_dotenv()
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not found in environment variables")
llm = ChatGoogleGenerativeAI(model="gemini-1.5-flash", temperature=0.3)
embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
search = DuckDuckGoSearchRun()

def extract_text_from_local_path(path):
    if path.endswith(".pdf"):
        with pdfplumber.open(path) as pdf:
            return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])
    elif path.endswith(".txt"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif path.endswith(".docx"):
        doc = DocxDocument(path)
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def router_agent(state):
    query = state.get("query", "")
    route_prompt = PromptTemplate.from_template(
        "Classify the query into one of [web, rag, llm]:\n\nQuery: {query}\n\nAnswer:"
    )
    route_result = (route_prompt | llm).invoke({"query": query}).content.lower()
    route = "llm"
    if "web" in route_result:
        route = "web"
    elif "rag" in route_result:
        route = "rag"
    return {**state, "route": route}

def web_agent(state):
    query = state["query"]
    try:
        result = search.run(query)
        return {**state, "content": result}
    except Exception as e:
        return {**state, "content": f"Web search failed: {str(e)}"}

def rag_agent(state):
    query = state["query"]
    retriever = state["retriever"]
    qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=retriever)
    answer = qa_chain.run(query)
    return {**state, "content": answer}

def llm_agent(state):
    query = state["query"]
    response = llm.invoke(query)
    return {**state, "content": response.content}

def summarizer_agent(state):
    content = state["content"]
    prompt = PromptTemplate.from_template("Summarize clearly and concisely:\n\n{content}")
    summary = (prompt | llm).invoke({"content": content}).content
    return {**state, "final": summary}

def run_langgraph(user_query, retriever):
    workflow = StateGraph(dict)
    workflow.set_entry_point("router")

    workflow.add_node("router", RunnableLambda(router_agent))
    workflow.add_node("web", RunnableLambda(web_agent))
    workflow.add_node("rag", RunnableLambda(rag_agent))
    workflow.add_node("llm", RunnableLambda(llm_agent))
    workflow.add_node("summarizer", RunnableLambda(summarizer_agent))

    def router_logic(state): return state["route"]
    workflow.add_conditional_edges("router", router_logic, {
        "web": "web",
        "rag": "rag",
        "llm": "llm"
    })

    for node in ["web", "rag", "llm"]:
        workflow.add_edge(node, "summarizer")

    workflow.set_finish_point("summarizer")
    app = workflow.compile()
    return app.invoke({"query": user_query, "retriever": retriever})["final"]

st.set_page_config(page_title="🔍 Agentic Research Assistant", layout="centered")
st.markdown("""
    <style>
    .main-title {
        text-align: center;
        font-size: 36px;
        font-weight: 700;
        margin-bottom: 10px;
        color: #3b82f6;
    }
    .subtitle {
        text-align: center;
        font-size: 18px;
        color: #6b7280;
    }
    .status-badge {
        display: inline-block;
        background-color: #e0f2fe;
        color: #0284c7;
        padding: 4px 12px;
        border-radius: 12px;
        font-size: 14px;
        margin-left: 10px;
    }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">Multi-Agent Research Assistant</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">LangGraph + Gemini + FAISS + Web Search + RAG</div>', unsafe_allow_html=True)
st.markdown("---")

retriever = None
documents_loaded = False
loaded_files = []

with st.spinner("Scanning 'data' folder for documents..."):
    if os.path.exists("data"):
        all_content = []
        for filename in os.listdir("data"):
            filepath = os.path.join("data", filename)
            if filename.lower().endswith(('.pdf', '.txt', '.docx')):
                content = extract_text_from_local_path(filepath)
                if content:
                    all_content.append(content)
                    loaded_files.append(filename)

        if all_content:
            chunks = text_splitter.create_documents(all_content)
            vectorstore = FAISS.from_documents(chunks, embeddings)
            retriever = vectorstore.as_retriever()
            documents_loaded = True
            st.success(f"Loaded {len(loaded_files)} documents.")
            with st.expander("View loaded files"):
                for f in loaded_files:
                    st.markdown(f"- {f}")
        else:
            st.warning("No valid files found in 'data' folder.")
    else:
        st.info(" 'data' folder does not exist.")

if not documents_loaded:
    st.info("Using fallback knowledge base.")
    st.markdown('<div class="status-badge">Fallback KB Active</div>', unsafe_allow_html=True)
    docs = [
        Document(page_content="LangGraph is a Python framework for agent workflows."),
        Document(page_content="Gemini 1.5 Flash is fast and great for summarization."),
    ]
    vectorstore = FAISS.from_documents(docs, embeddings)
    retriever = vectorstore.as_retriever()

st.markdown("---")

# User Query Input
with st.form(key="query_form"):
    st.markdown("### Ask a question:")
    query = st.text_input("", placeholder="e.g. What is LangGraph?", label_visibility="collapsed")
    col1, col2 = st.columns([0.7, 0.3])
    with col2:
        submit = st.form_submit_button("Submit")

if submit:
    if not query.strip():
        st.warning("Please enter a question.")
    else:
        with st.spinner("analysing..."):
            try:
                answer = run_langgraph(query, retriever)
                st.success("Answer generated successfully!")
                st.markdown("### Answer:")
                st.write(answer)
            except Exception as e:
                st.error(f"Error: {str(e)}")
